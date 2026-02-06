import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk  # 仅用于 Progressbar
from tkinter.scrolledtext import ScrolledText
import tkinter.font as tkfont
import threading
import os
import pandas as pd
from docx import Document
import datetime
import numbers


class EmailMergeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("邮件批量合并工具-pythonfun作品")
        # 正方形窗口 & 原始风格
        self.root.geometry("680x380")
        self.root.resizable(False, False)

        # 变量
        self.template_var = tk.StringVar()
        self.data_var = tk.StringVar()
        self.status_var = tk.StringVar(value="就绪")
        self.single_file_var = tk.BooleanVar(value=False)  # 勾选则单文件合并
        self.total_rows = 0

        # 只设置界面字体（不影响生成的 Word 文档）
        self._apply_ui_font(("宋体", "SimSun"), 11)

        self._build_ui()
        self._wire_events()

    # ---------- 仅改 UI 默认字体 ----------
    def _apply_ui_font(self, family_candidates=("宋体", "SimSun"), size=12):
        families = set(tkfont.families())
        for cand in family_candidates:
            if cand in families:
                family = cand
                break
        else:
            family = tkfont.nametofont("TkDefaultFont").cget("family")

        for name in ("TkDefaultFont", "TkTextFont", "TkFixedFont", "TkMenuFont", "TkHeadingFont"):
            try:
                f = tkfont.nametofont(name)
                f.config(family=family, size=size)
            except tk.TclError:
                pass

    # ---------- UI ----------
    def _build_ui(self):
        main = tk.Frame(self.root, padx=12, pady=12)
        main.pack(fill=tk.BOTH, expand=True)

        # 第一行：模板 / 数据（路径在左、按钮在右）
        row = tk.Frame(main)
        row.grid(row=0, column=0, sticky="ew", pady=(0, 12))
        for c in range(8):
            row.grid_columnconfigure(c, weight=0)
        row.grid_columnconfigure(1, weight=1)  # 模板路径伸缩
        row.grid_columnconfigure(5, weight=1)  # 数据路径伸缩

        tk.Label(row, text="模板：").grid(row=0, column=0, sticky="w", padx=(0, 6))
        self.template_entry = tk.Entry(row, textvariable=self.template_var)
        self.template_entry.grid(row=0, column=1, sticky="ew", padx=(0, 8))
        self.btn_template = tk.Button(row, text="选择模板", command=self._pick_template)
        self.btn_template.grid(row=0, column=2, sticky="e")

        tk.Label(row, text="数据表：").grid(row=0, column=4, sticky="w", padx=(18, 6))
        self.data_entry = tk.Entry(row, textvariable=self.data_var)
        self.data_entry.grid(row=0, column=5, sticky="ew", padx=(0, 8))
        self.btn_data = tk.Button(row, text="选择数据", command=self._pick_data)
        self.btn_data.grid(row=0, column=6, sticky="e")

        # 同一行左右分栏：左侧选项，右侧按钮
        controls = tk.Frame(main)
        controls.grid(row=2, column=0, sticky="ew", pady=(0, 12))
        controls.grid_columnconfigure(0, weight=1)  # 左列可伸展，按钮始终靠右

        # 左侧：选项
        opts = tk.Frame(controls)
        opts.grid(row=0, column=0, sticky="ew")
        opts.grid_columnconfigure(2, weight=1)  # 让第2列（combobox）可拉伸

        self.ck_single = tk.Checkbutton(opts, text="单文件", variable=self.single_file_var)
        self.ck_single.grid(row=0, column=0, sticky="w")

        self.field_prefix_label = tk.Label(opts, text="文件名前缀：")
        self.field_prefix_label.grid(row=0, column=1, sticky="w", padx=(12, 6))

        self.field_prefix_combobox = ttk.Combobox(opts, state="readonly")
        self.field_prefix_combobox.grid(row=0, column=2, sticky="ew", padx=(0, 12))

        # 右侧：按钮
        ops = tk.Frame(controls)
        ops.grid(row=0, column=1, sticky="e")
        self.btn_merge = tk.Button(ops, text="邮件合并", command=self._start_merge)
        self.btn_reset = tk.Button(ops, text="重置程序", command=self._reset)
        self.btn_exit = tk.Button(ops, text="退出程序", command=self.root.destroy)
        self.btn_merge.grid(row=0, column=0, padx=(0, 8))
        self.btn_reset.grid(row=0, column=1, padx=(0, 8))
        self.btn_exit.grid(row=0, column=2)

        # 进度条（ttk Progressbar）
        prog = tk.Frame(main)
        prog.grid(row=3, column=0, sticky="ew", pady=(0, 12))
        prog.grid_columnconfigure(0, weight=1)
        self.progress = ttk.Progressbar(prog, mode="determinate", maximum=100)
        self.progress.grid(row=0, column=0, sticky="ew")
        self.progress_pct = tk.Label(prog, width=6, anchor="e")
        self.progress_pct.grid(row=0, column=1, padx=(6, 0))

        # 占位
        filler = tk.Frame(main)
        filler.grid(row=4, column=0, sticky="nsew")
        filler.grid_rowconfigure(0, weight=0)  # 不扩展
        filler.grid_columnconfigure(0, weight=1)
        tk.Label(
            filler,
            text="模板中占位符使用 {{字段名}}，字段名与数据表头一致；不要把单个占位符分段加粗/变色。"
        ).grid(sticky="nw")

        # 状态区（日志 + 单行状态）
        status = tk.Frame(main)
        status.grid(row=5, column=0, sticky="nsew")
        main.grid_rowconfigure(5, weight=1)  # 让日志区可扩展
        status.grid_columnconfigure(0, weight=1)
        status.grid_rowconfigure(0, weight=1)

        # 日志（可滚动、不可编辑）
        self.status_log = ScrolledText(status, height=6, state="disabled")
        self.status_log.grid(row=0, column=0, sticky="nsew")
        self.status_log.tag_config("ok", foreground="#2e7d32")
        self.status_log.tag_config("err", foreground="#c62828")
        self.status_log.tag_config("info", foreground="#555555")

        # 单行状态 Label（即时提示）
        self.status_label = tk.Label(
            status, textvariable=self.status_var, anchor="w",
            relief="groove", padx=8, pady=4
        )
        self.status_label.grid(row=1, column=0, sticky="ew")

    def _wire_events(self):
        self.template_var.trace_add("write", lambda *_: self._update_start_state())
        self.data_var.trace_add("write", lambda *_: self._update_start_state())
        self._update_start_state()

    # ---------- 事件/校验 ----------
    def _pick_template(self):
        path = filedialog.askopenfilename(
            title="选择 Word 模板（.docx）",
            filetypes=[("Word 文件", "*.docx")],
            initialdir=os.getcwd()
        )
        if not path:
            return
        if not path.lower().endswith(".docx"):
            messagebox.showerror("文件类型不正确", "模板必须为 .docx。")
            return
        self.template_var.set(path)
        self._set_status(f"已选择模板：{os.path.basename(path)}")
        self._log(f"已选择模板：{os.path.basename(path)}")

    def _pick_data(self):
        path = filedialog.askopenfilename(
            title="选择数据文件（.xlsx/.xls）",
            filetypes=[("Excel 文件", "*.xlsx;*.xls")],
            initialdir=os.getcwd()
        )
        if not path:
            return
        if os.path.splitext(path)[1].lower() not in (".xlsx", ".xls"):
            messagebox.showerror("文件类型不正确", "数据文件必须为 .xlsx 或 .xls。")
            return
        self.data_var.set(path)

        # 读取数据列名填充 combobox
        try:
            df = pd.read_excel(path)  # 只为取列名
            columns = list(df.columns)
            self.field_prefix_combobox["values"] = columns
            self._set_status(f"已选择数据：{os.path.basename(path)}")
            self._log(f"已选择数据：{os.path.basename(path)}（列：{', '.join(map(str, columns))}）")
        except Exception as e:
            messagebox.showerror("错误", f"读取数据失败：{e}")
            self._log(f"读取数据失败：{e}", ok=False)

    def _update_start_state(self):
        t_ok = self.template_var.get().lower().endswith(".docx")
        d_ok = self.data_var.get().lower().endswith((".xlsx", ".xls"))
        self.btn_merge.configure(state=(tk.NORMAL if (t_ok and d_ok) else tk.DISABLED))

    # ---------- 合并逻辑（线程） ----------
    def _start_merge(self):
        if not self.template_var.get() or not self.data_var.get():
            messagebox.showwarning("缺少文件", "请先选择模板与数据文件。")
            return
        if not self.template_var.get().lower().endswith(".docx"):
            messagebox.showerror("文件类型不正确", "模板必须为 .docx。")
            return
        if not self.data_var.get().lower().endswith((".xlsx", ".xls")):
            messagebox.showerror("文件类型不正确", "数据文件必须为 .xlsx 或 .xls。")
            return

        self._set_busy(True)
        self._set_status("开始处理……")
        self.progress["value"] = 0
        self.progress_pct.config(text="0%")
        self._ui(lambda: self._log("开始处理……"))

        threading.Thread(target=self._merge_worker, daemon=True).start()

    # 文件名清洗（可能返回空串，用于兜底）
    def _safe_name(self, s):
        s = "" if s is None else str(s).strip()
        bad = '<>:"/\\|?*'
        cleaned = "".join(ch for ch in s if ch not in bad).strip()
        return cleaned

    def _merge_worker(self):
        try:
            # 读数据
            data_path = self.data_var.get()
            try:
                if data_path.lower().endswith(".xls"):
                    try:
                        df = pd.read_excel(data_path, engine="xlrd")
                    except Exception:
                        df = pd.read_excel(data_path)  # 回退
                else:
                    df = pd.read_excel(data_path, engine="openpyxl")
            except Exception as read_e:
                raise RuntimeError(f"读取数据失败：{read_e}")

            if df.empty:
                raise ValueError("数据文件为空。")

            template_path = self.template_var.get()
            out_dir = os.path.join(os.path.dirname(template_path), "合并输出")
            os.makedirs(out_dir, exist_ok=True)

            total = len(df)
            single = self.single_file_var.get()

            # 尝试单文件合并（docxcompose）
            composer = None
            if single:
                try:
                    from docxcompose.composer import Composer
                    composer_available = True
                except Exception:
                    composer_available = False
                    self._ui(lambda: self._set_status("未检测到 docxcompose，将改为多文件输出。"))
                    self._ui(lambda: self._log("未检测到 docxcompose，将改为多文件输出。"))
                if composer_available:
                    first_doc = None
                else:
                    single = False  # 降级为多文件

            # 逐行生成
            for i, (_, row) in enumerate(df.iterrows(), start=1):
                # 修改这里：更全面地处理日期类型数据
                mapping = {}
                for k, v in row.items():
                    if pd.isna(v):
                        mapping[str(k)] = ""
                    elif isinstance(v, (pd.Timestamp, datetime.datetime)):
                        # 将日期格式化为字符串
                        mapping[str(k)] = v.strftime('%Y-%m-%d')
                    elif isinstance(v, datetime.date):
                        # 处理 date 类型
                        mapping[str(k)] = v.strftime('%Y-%m-%d')
                    elif isinstance(v, numbers.Number):
                        # 检查是否是 Excel 日期序列号
                        try:
                            # Excel 日期序列号是从 1900-01-01 开始的天数
                            # 但要注意 Excel 错误地将 1900 年视为闰年
                            if 0 < v < 100000:  # 合理的日期序列号范围
                                # 转换为日期
                                excel_epoch = datetime.datetime(1899, 12, 30)  # Excel 的起始日期
                                date_value = excel_epoch + datetime.timedelta(days=v)
                                mapping[str(k)] = date_value.strftime('%Y-%m-%d')
                            else:
                                mapping[str(k)] = str(v)
                        except:
                            mapping[str(k)] = str(v)
                    else:
                        mapping[str(k)] = str(v)

                doc = Document(template_path)
                self._replace_placeholders(doc, mapping)

                if single:
                    if i == 1:
                        first_doc = doc
                        composer = Composer(first_doc)
                    else:
                        first_doc.add_page_break()
                        composer.append(doc)
                else:
                    # 新：按所选列在当前行的值命名（空值兜底）
                    col = self.field_prefix_combobox.get().strip()
                    cell_val = row[col] if (col in row and not pd.isna(row[col])) else ""
                    prefix_val = self._safe_name(cell_val)
                    out_name = f"{prefix_val}_邮件_{i:03d}.docx" if prefix_val else f"邮件_{i:03d}.docx"
                    out_path = os.path.join(out_dir, out_name)
                    doc.save(out_path)

                # 进度 + 日志
                pct = int(i / total * 100)
                self._ui(lambda p=pct, i=i: (
                    self.progress.config(value=p),
                    self.progress_pct.config(text=f"{p}%"),
                    self._set_status(f"处理 {i}/{total}"),
                    self._log(f"处理 {i}/{total}")
                ))

            # 保存单文件
            if single and composer is not None:
                # 合并文件命名：用所选列第1行的值；为空则默认"邮件_合并版.docx"
                col = self.field_prefix_combobox.get().strip()
                first_val = df.iloc[0][col] if (
                            col in df.columns and not df.empty and not pd.isna(df.iloc[0][col])) else ""
                merged_name = f"{self._safe_name(first_val)}_合并版.docx" if self._safe_name(first_val) else "邮件_合并版.docx"
                merged_path = os.path.join(out_dir, merged_name)
                composer.save(merged_path)
                self._ui(lambda: self._set_status(f"完成！单文件已生成：{merged_path}"))
                self._ui(lambda: self._log(f"合并完成：{merged_path}", ok=True))
                self._ui(lambda: messagebox.showinfo("完成", f"合并完成！单文件已输出到：\n{merged_path}"))
            else:
                self._ui(lambda: self._set_status(f"完成！多文件已生成至目录：{out_dir}"))
                self._ui(lambda: self._log(f"合并完成：多文件输出到 {out_dir}", ok=True))
                self._ui(lambda: messagebox.showinfo("完成", f"合并完成！多文件已输出到：\n{out_dir}"))

        except Exception as e:
            err = str(e)  # 修复：捕获错误文本用于异步 UI
            self._ui(lambda err=err: messagebox.showerror("错误", f"合并失败：{err}"))
            self._ui(lambda err=err: self._set_status(f"错误：{err}"))
            self._ui(lambda err=err: self._log(f"合并失败：{err}", ok=False))
        finally:
            self._ui(lambda: self._set_busy(False))

    # ---------- 占位符替换 ----------
    def _replace_placeholders(self, doc: Document, mapping: dict):
        """保留原文样式，同时替换占位符（占位符应在同一 run 内）"""

        def replace_in_paragraph(p):
            if not p.runs:
                return
            for run in p.runs:
                txt = run.text
                replaced = False
                for k, v in mapping.items():
                    # 支持 {{key}} 与 {{ key }}
                    for pat in (f"{{{{{k}}}}}", f"{{{{ {k} }}}}"):
                        if pat in txt:
                            txt = txt.replace(pat, v)
                            replaced = True
                if replaced:
                    run.text = txt  # 仅原位替换，不清段、不新建 run

        def replace_in_table(table):
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)

        # 正文
        for p in doc.paragraphs:
            replace_in_paragraph(p)

        # 表格
        for t in doc.tables:
            replace_in_table(t)

        # 页眉/页脚（默认/首页/偶数页）
        for sec in doc.sections:
            for attr in ("header", "first_page_header", "even_page_header"):
                hdr = getattr(sec, attr, None)
                if hdr:
                    for p in hdr.paragraphs:
                        replace_in_paragraph(p)
                    for t in hdr.tables:
                        replace_in_table(t)
            for attr in ("footer", "first_page_footer", "even_page_footer"):
                ftr = getattr(sec, attr, None)
                if ftr:
                    for p in ftr.paragraphs:
                        replace_in_paragraph(p)
                    for t in ftr.tables:
                        replace_in_table(t)

    # ---------- 工具 ----------
    def _log(self, text: str, ok=None):

        # ok=True &#8658; &#10004;，ok=False &#8658; &#10006;，None &#8658; &#8226;
        prefix = "\u2714 " if ok is True else ("\u2716 " if ok is False else "\u2022")
        tag = "ok" if ok is True else ("err" if ok is False else "info")
        self.status_log.configure(state="normal")
        self.status_log.insert("end", prefix + text + "\n", tag)
        self.status_log.see("end")
        self.status_log.configure(state="disabled")

    def _set_status(self, text):
        self.status_var.set(text)

    def _set_busy(self, busy: bool):
        state = (tk.DISABLED if busy else tk.NORMAL)
        # 只对需要交互的控件改状态；进度条不动
        for w in (
                self.btn_template, self.btn_data, self.btn_merge, self.btn_reset, self.btn_exit,
                self.template_entry, self.data_entry, self.ck_single, self.field_prefix_combobox
        ):
            w.configure(state=state)
        self.root.config(cursor="watch" if busy else "")

    def _reset(self):
        self.template_var.set("")
        self.data_var.set("")
        self.single_file_var.set(False)
        self.progress["value"] = 0
        self.progress_pct.config(text="0%")
        self._set_status("已重置")
        self._log("已重置")
        self.template_entry.focus_set()

    def _ui(self, fn):
        self.root.after(0, fn)


# 入口
if __name__ == "__main__":
    root = tk.Tk()
    EmailMergeApp(root)
    root.mainloop()