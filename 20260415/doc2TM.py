import os
import re
import sys
import configparser
import spacy
from docx import Document
import win32com.client as win32


# ================= 1. 读取配置文件 =================
def load_config():
    config = configparser.ConfigParser()
    config.optionxform = str  # 强制保留键的大小写，防止大写字母变小写

    # 确定配置文件的路径 (兼容打包成 EXE 后的运行路径)
    if getattr(sys, 'frozen', False):
        application_path = os.path.dirname(sys.executable)
    else:
        application_path = os.path.dirname(os.path.abspath(__file__))

    config_file = os.path.join(application_path, "doc2TM.ini")

    if not os.path.exists(config_file):
        print(f"【错误】找不到配置文件: {config_file}")
        print("请确保 config.ini 与本程序放在同一文件夹下！")
        os.system("pause")
        sys.exit(1)

    # 读取文件，使用 utf-8-sig 兼容带 BOM 的 UTF-8 编码
    config.read(config_file, encoding='utf-8-sig')

    # 安全地读取 LogoScanNum（如果没填或填了非数字，默认用20）
    try:
        logo_scan_num = int(config.get("Settings", "LogoScanNum", fallback=20))
    except ValueError:
        print("【警告】LogoScanNum 配置不正确，已自动重置为默认值 20")
        logo_scan_num = 20

    # 解析配置
    settings = {
        "NAME_PLACEHOLDER": config.get("Settings", "NamePlaceholder", fallback="某某"),
        "ORG_PLACEHOLDER": config.get("Settings", "OrgPlaceholder", fallback="某单位"),
        "LOGO_SCAN_NUM": logo_scan_num
    }

    whitelist_str = config.get("WhiteList", "Items", fallback="")
    whitelist = [item.strip() for item in whitelist_str.split(",") if item.strip()]

    replacements = dict(config.items("Replacements")) if config.has_section("Replacements") else {}

    return settings, whitelist, replacements


# 加载配置
SETTINGS, WHITE_LIST, CUSTOM_REPLACEMENTS = load_config()

# ================= 2. 加载 AI 模型 =================
print("正在加载 AI 模型，请稍候...")
try:
    nlp = spacy.load("zh_core_web_sm")
except Exception as e:
    print(f"模型加载失败: {e}")
    print("【错误】缺少 zh_core_web_sm 模型！")
    os.system("pause")
    sys.exit(1)


# ================= 3. 核心处理逻辑 =================
def smart_anonymize(text):
    if not text or not text.strip():
        return text

    # ====================================================
    # 新增：第一步 —— 开启“绝对保护罩”
    # 将文本中出现的白名单词汇，临时替换成特殊占位符 (如 [[WP_0]])
    # ====================================================
    protected_items = {}
    for i, word in enumerate(WHITE_LIST):
        # 忽略大小写进行查找和保护 (比如保护 cpu 也会保护 CPU)
        pattern = re.compile(re.escape(word), re.IGNORECASE)
        # 查找文本中是否包含该白名单词
        matches = set(pattern.findall(text))
        for match_str in matches:
            placeholder = f"[[WP_{i}]]"
            protected_items[placeholder] = match_str
            text = text.replace(match_str, placeholder)

    # ====================================================
    # 第二步 —— 正常的脱密逻辑 (此时白名单词已被隐藏，绝对安全)
    # ====================================================
    # 1. 关键词替换
    sorted_keywords = sorted(CUSTOM_REPLACEMENTS.keys(), key=len, reverse=True)
    for old in sorted_keywords:
        text = text.replace(old, CUSTOM_REPLACEMENTS[old])

    # 2. 正则匹配单位 (此时 "结构布局" 已经变成了 "结构布[[WP_x]]"，不会再被 "局" 误杀)
    unit_pattern = r"([\u4e00-\u9fa5]{2,}(?:有限公司|股份公司|公司|研究所|集团|中心|委员会|项目组|厂|局))"
    text = re.sub(unit_pattern, SETTINGS["ORG_PLACEHOLDER"], text)

    # 3. AI 识别
    doc = nlp(text)
    entities = sorted(doc.ents, key=lambda x: len(x.text), reverse=True)
    for ent in entities:
        # 只要不是特殊占位符，就正常替换
        if not ent.text.startswith("[[WP_"):
            if ent.label_ == "PERSON":
                text = text.replace(ent.text, SETTINGS["NAME_PLACEHOLDER"])
            elif ent.label_ == "ORG" and ent.text not in CUSTOM_REPLACEMENTS.values():
                text = text.replace(ent.text, SETTINGS["ORG_PLACEHOLDER"])

    # ====================================================
    # 新增：第三步 —— 解除保护，恢复白名单词汇
    # ====================================================
    for placeholder, original_word in protected_items.items():
        text = text.replace(placeholder, original_word)

    return text


def clean_images_via_word_engine(input_path):
    print(f"  -> 正在启动 Word 内核清理图片...")
    word = win32.DispatchEx('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0

    try:
        doc = word.Documents.Open(os.path.abspath(input_path))

        # 扫荡页眉页脚的所有图形和嵌入式图片
        for section in doc.Sections:
            for hf_type in [1, 2, 3]:
                header = section.Headers(hf_type)
                if header.Exists:
                    for i in range(header.Shapes.Count, 0, -1): header.Shapes(i).Delete()
                    for i in range(header.Range.InlineShapes.Count, 0, -1): header.Range.InlineShapes(i).Delete()

                footer = section.Footers(hf_type)
                if footer.Exists:
                    for i in range(footer.Shapes.Count, 0, -1): footer.Shapes(i).Delete()
                    for i in range(footer.Range.InlineShapes.Count, 0, -1): footer.Range.InlineShapes(i).Delete()

        # 动态获取扫描段落数
        scan_num = SETTINGS["LOGO_SCAN_NUM"]
        check_limit = min(scan_num, doc.Paragraphs.Count)

        if check_limit > 0:
            limit_end = doc.Paragraphs(check_limit).Range.End
            for i in range(1, check_limit + 1):
                rng = doc.Paragraphs(i).Range
                for j in range(rng.InlineShapes.Count, 0, -1): rng.InlineShapes(j).Delete()

            for i in range(doc.Shapes.Count, 0, -1):
                try:
                    if doc.Shapes(i).Anchor.Start <= limit_end:
                        doc.Shapes(i).Delete()
                except Exception:
                    pass

        temp_docx = os.path.abspath(input_path) + "_temp.docx"
        doc.SaveAs(temp_docx, FileFormat=16)
        doc.Saved = True
        doc.Close()
        return temp_docx

    except Exception as e:
        print(f"    Word 内核处理失败: {e}")
        raise e
    finally:
        word.Quit(0)


def process_document(input_file, output_file):
    print(f"\n正在处理: {os.path.basename(input_file)}")
    temp_working_file = clean_images_via_word_engine(input_file)

    try:
        print(f"  -> 正在进行文字脱密...")
        doc = Document(temp_working_file)

        # ================== 替换这里的函数 ==================
        def handle_paragraphs(paras):
            for p in paras:
                original_text = p.text
                if not original_text.strip():
                    continue

                new_text = smart_anonymize(original_text)
                if new_text != original_text:
                    if p.runs:
                        p.runs[0].text = new_text
                        for i in range(1, len(p.runs)):
                            p.runs[i].text = ""

        # ====================================================

        handle_paragraphs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    handle_paragraphs(cell.paragraphs)

        for section in doc.sections:
            if section.header:
                handle_paragraphs(section.header.paragraphs)
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells: handle_paragraphs(cell.paragraphs)
            if section.footer:
                handle_paragraphs(section.footer.paragraphs)

        output_file_docx = output_file if output_file.endswith(".docx") else output_file + "x"
        doc.save(output_file_docx)
        print(f"  √ 成功完成: {os.path.basename(output_file_docx)}")

    finally:
        if os.path.exists(temp_working_file):
            os.remove(temp_working_file)


def batch_process_documents():
    if getattr(sys, 'frozen', False):
        current_dir = os.path.dirname(sys.executable)
    else:
        current_dir = os.getcwd()

    extensions = ('.doc', '.docx')
    print(f"====== 开始静默批量脱密处理 ======")

    file_count = 0
    for file_name in os.listdir(current_dir):
        if (file_name.lower().endswith(extensions) and
                not file_name.startswith("~$") and
                not file_name.rsplit('.', 1)[0].endswith('_TM') and
                not file_name.endswith('_temp.docx')):

            input_path = os.path.join(current_dir, file_name)
            name_part, extension = os.path.splitext(file_name)
            output_name = f"{name_part}_TM{extension}"
            output_path = os.path.join(current_dir, output_name)

            try:
                process_document(input_path, output_path)
                file_count += 1
            except Exception as e:
                print(f"  × 处理出错 {file_name}: {e}")

    print("==============================")
    print(f"所有文件处理完毕！共处理了 {file_count} 个文件。")
    os.system("pause")


if __name__ == "__main__":
    batch_process_documents()